# -*- coding: utf-8 -*-
"""生成面试备战手册 Word 文档。
新增/修改的内容用红色标出（用户已打印纸质版，方便对照红笔标注）。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 默认字体
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(10.5)
style.font.color.rgb = BLACK


def h1(text, red=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RED if red else BLACK
    p.space_before = Pt(14)
    return p


def h2(text, red=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12.5)
    r.font.color.rgb = RED if red else BLACK
    p.space_before = Pt(8)
    return p


def body(text, red=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = RED if red else BLACK
    r.bold = bold
    return p


def table(headers, rows, red_cols=()):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9.5)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    if ci in red_cols:
                        r.font.color.rgb = RED
    return t


# ═══════════════════════════════════════════════════════════
# 封面标题
# ═══════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('8个AI应用项目 — 面试备战手册')
r.bold = True
r.font.size = Pt(22)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('技术栈 → 架构设计 → 核心难点 → Bug解决记录\n（红色 = 2026-08-13 最新更新，对照纸质版用红笔标注）')
r.font.size = Pt(10)
r.font.color.rgb = RED

# ═══════════════════════════════════════════════════════════
# 一、技术框架
# ═══════════════════════════════════════════════════════════
h1('一、技术框架')
body('前端: Vue3 + TypeScript + NaiveUI + Pinia + Vue Router + Vite')
body('后端: Python FastAPI + SQLAlchemy 2.0 (async) + JWT认证 + SSE流式')
body('AI: LangChain + DeepSeek + ChromaDB 向量检索 + RAG')
body('数据: SQLite WAL（开发），8项目端口独立互不冲突')
body('联动: X-Internal-Call 数据飞轮 — ①⑤④→⑥数据中心→⑧微调→推理服务', red=True, bold=True)

# ═══════════════════════════════════════════════════════════
# 二、8个项目概览
# ═══════════════════════════════════════════════════════════
h1('二、8个项目概览')
table(
    ['#', '项目', '定位', '核心功能', '端口'],
    [
        ['①', 'RAG智能客服系统', 'C端', '知识库问答、转人工、引用溯源', '8101:3001'],
        ['②', 'AI自媒体内容助手', 'C端', '爆款标题/脚本/图文生成，5平台适配，多AI供应商回退', '8202:3002'],
        ['③', 'AI短视频脚本工坊', 'C端', '分镜表+口播稿+TTS语音+字幕导出，多AI供应商回退', '8000:3000'],
        ['④', 'AI素材管理平台', 'B端', '素材上传/AI打标/图搜图/Unsplash+Pexels外部图库', '8400:3004'],
        ['⑤', 'AI销售培训系统', 'B端', 'AI扮演客户/5维打分/进步曲线', '8505:3005'],
        ['⑥', 'AI数据中心平台', '中台', '采集→清洗→标注→版本对比→质量报告', '8606:3006'],
        ['⑦', 'MCP多智能体协作平台', '中台', '任务自动拆解+4种编排模式+SSE监控', '8707:3007'],
        ['⑧', 'AI模型微调训练平台', '中台', '微调任务/Loss曲线/RLHF/A-B对比/部署', '8808:3008'],
    ],
    red_cols=(3,),
)
body('【红色标注】②③新增多AI供应商回退（DeepSeek失败→智谱GLM→通义千问）', red=True)
body('【红色标注】④新增Unsplash+Pexels外部图库API（搜索+一键导入+AI打标）', red=True)

# ═══════════════════════════════════════════════════════════
# 三、核心难点（面试重点）
# ═══════════════════════════════════════════════════════════
h1('三、核心难点 & Bug 解决（面试重点）')

h2('难点 1：LangChain Prompt 花括号冲突')
body('现象：KeyError: missing variables {功能, 场景, 效果}')
body('根因：prompts.py 中 {xxx} 被 ChatPromptTemplate 当变量匹配')
body('解决：双写花括号转义 {{功能}}，当作普通文本')
body('影响：①②③的 prompts.py 全部修复')

h2('难点 2：端口冲突 → 连环错')
body('现象：项目②回复"LLM未配置"、检索到①的旧数据')
body('排查：curl 8000 端口返回 "RAG智能客服系统" → 端口被①占了')
body('解决：8项目分配独立端口（8101/8202/8000…）')

h2('难点 3：SQLite → WAL → PostgreSQL 性能三阶段')
body('50并发压测：SQLite标准 62%失败 → WAL+5项PRAGMA 30% → PostgreSQL 6%')
body('PRAGMA: journal_mode=WAL, synchronous=NORMAL, cache_size=-8000, busy_timeout=5000, foreign_keys=ON')

h2('难点 4：Windows 上传中文文件名乱码（GBK→Latin-1→UTF-8）')
body('根因：浏览器发GBK字节，FastAPI按Latin-1解码 → 乱码')
body('解决：_fix_filename() 用 encode("latin-1").decode("gbk") 修复')

h2('难点 5：上传后一直 pending → 异步改同步的架构权衡')
body('旧方案 asyncio.create_task 后台处理，重启即丢。新方案同步处理，用户看到的一定是 completed')
body('面试要点：能同步就不异步，消息队列是最后手段。50MB内文档处理<5秒，同步完全可行')

h2('难点 6：DOCX/PDF 无法预览 → LangChain Loader')
body('docx 是 ZIP 压缩的 XML，不能 open().read()。用 Docx2txtLoader + PyPDFLoader 提取文本')

h2('难点 7：change-me 占位符 → 全项目密钥分离')
body('审查发现 8 个 .env 文件硬编码真实 DeepSeek Key。替换为占位符 + .env.example 模板 + .gitignore 兜底')
body('面试要点：最小权限原则——开发环境也用占位符，部署时 CI/CD 注入')

h2('难点 8：seed_admin_user 密码不同步 → 改 .env 后登录失败')
body('根因：代码只在用户不存在时创建，.env 改了密码但 DB 还是旧哈希')
body('解决：启动时 verify_password 比对，不一致自动重哈希同步。配置≠数据库')

h2('难点 9：文件端点无认证 → 任意下载（安全审查）')
body('④的 5 个 CRUD 端点中 4 个有认证，唯独文件下载漏了')
body('面试要点：逐行检查每个端点的认证依赖——需不需要登录？有没有 Depends？有没有所有权检查？')

h2('难点 10：forgot-password 用户枚举漏洞')
body('hint 字段区分"用户存在/未注册"→ 攻击者可批量探测用户名')
body('解决：统一返回"如果该账号存在，密码重置链接已发送"')

h2('难点 11：跨项目推送全部 401 → X-Internal-Call 内部认证', red=True)
body('现象：①⑤⑧推送数据到⑥全部被 401 拦截，数据飞轮是断的', red=True)
body('根因：⑥要求 JWT，推送方只发 X-Internal-Call: true——后端根本没实现该头的验证', red=True)
body('解决：⑥新增 INTERNAL_CALL_SECRET 共享密钥 + internal_call_or_user 依赖', red=True)
body('面试说法：微服务内部认证两条路——mTLS（重）或共享密钥头（轻）。8项目规模用共享密钥+独立数据库隔离', red=True)

h2('难点 12：SSE 执行 ID 在 flush 前取用 → 4种编排模式全崩（⑦）')
body('根因：db.add(ex) 后立刻 str(ex.id)——UUID 默认值在 flush/INSERT 时才求值，此时是 None')
body('解决：先收集对象 → await db.flush() → 再取 ID。用内存SQLite实测验证了三个时机')

h2('难点 13：Windows GBK 控制台 emoji 崩溃')
body('现象：CI全绿但Windows本地启动即崩——print("⚠️") 在GBK编码下 UnicodeEncodeError')
body('解决：emoji 换 ASCII + 启动脚本 PYTHONIOENCODING=utf-8。print 调试信息不要用 emoji')

h2('难点 14：跨项目认证密钥硬编码 "true" → 静默 401', red=True)
body('现象：⑦⑧调用⑥的接口全部401，但代码看起来"认证了"', red=True)
body('根因：多处代码发 X-Internal-Call: "true" 字面量，⑥验证的是共享密钥。⑧的Few-shot层因此永远0样本', red=True)
body('解决：统一为真实密钥 + 双URL回退（Docker服务名→localhost）', red=True)
body('教训：内部认证密钥散落多文件容易漂移，应抽常量从配置读取', red=True)

h2('难点 15：上传标签静默丢失（缺 Form 注解）', red=True)
body('现象：④上传时手动填的标签永远存不上', red=True)
body('根因：tags 参数缺 Form(None) 注解——FastAPI 当 query 参数处理，FormData 里的标签被忽略', red=True)
body('解决：补注解 + 扩展名白名单（11种类型）。混合 multipart 上传最容易漏注解', red=True)

# ═══════════════════════════════════════════════════════════
# 四、架构亮点（面试加分）
# ═══════════════════════════════════════════════════════════
h1('四、架构亮点（面试加分）')

h2('SSE 流式响应')
body('后端 StreamingResponse + async generator 逐 token 产出；5种事件：thinking/retrieving/token/sources/done')
body('前端原生 fetch + ReadableStream 解析，AbortController 取消；三层回退：向量搜索→关键词→mock')

h2('多 AI 供应商自动回退（②③）', red=True)
body('DeepSeek 失败 → 智谱 GLM-4-Flash → 通义 Qwen-Turbo；OpenAI 兼容接口统一封装', red=True)
body('get_llm_with_failover() 三级优先：微调代理 → DeepSeek → 多供应商', red=True)
body('面试说法：多供应商容灾回退，单一 API 故障不影响服务可用性', red=True)

h2('外部图库 API 接入（④）', red=True)
body('Unsplash（图片）+ Pexels（照片+视频）+ Lorem Picsum（免Key）三图库；未配Key时优雅降级', red=True)
body('前端"🌍外部图库"：4数据源切换+搜索+一键导入素材库（自动AI打标）', red=True)

h2('数据飞轮（8项目联动）')
body('①客服对话 ⑤话术训练 ④素材标签 → ⑥数据中心（清洗+AI标注+版本化）→ ⑧模型微调 → 推理服务')
body('⑦运营引擎通过 action_registry 编排调用全部系统')
body('实测闭环：⑧从⑥导入数据集创建微调任务成功（dataset_items>0）', red=True)

h2('Smart Proxy 三层推理（⑧）')
body('意图路由（8领域关键词）→ 训练数据 Few-shot 检索 → LRU 缓存（MD5+TTL）')
body('实测：第一次调用 cached:False，第二次 cached:True（缓存命中）', red=True)

# ═══════════════════════════════════════════════════════════
# 五、面试常见追问预案
# ═══════════════════════════════════════════════════════════
h1('五、面试常见追问预案')

h2('Q: 为什么选 LangChain 而不是直接调 API？')
body('统一 Prompt 模板、Chain 组装、向量检索集成。8项目共享同一套 RAG pipeline，新增项目只需换 Prompt 和 sample-data')

h2('Q: 遇到的最难 bug 是什么？')
body('1. GBK 编码乱码链——表象是文件名乱码，排查了数据库/前端/HTTP header 才发现是 Windows→FastAPI 多层编码错误')
body('2. 端口冲突连环错——表象是 LLM 不工作，最后发现旧进程占了端口')

h2('Q: 为什么 SQLite 换 PostgreSQL？')
body('50并发 SQLite 锁竞争 62% 失败，WAL 改善到 30% 仍有瓶颈，PostgreSQL 真正并发写入降至 6%')

h2('Q: 为什么上传文件有时要等好久？')
body('旧版后台任务重启即丢。新版同步处理——用户看到的一定是 completed。能同步就不异步')

h2('Q: 暗色模式怎么实现？')
body('双层方案：NaiveUI darkTheme 负责组件 + [data-theme="dark"] CSS 变量 + Pinia 持久化用户偏好')

h2('Q: 8个项目如何联动？', red=True)
body('X-Internal-Call 共享密钥内部认证 + 数据飞轮：①⑤④推送数据→⑥清洗标注→⑧微调→推理服务；⑦编排调度', red=True)
body('实测验证：⑤结束训练会话后⑥自动收到数据；⑧从⑥导入数据集创建微调任务成功', red=True)

# ═══════════════════════════════════════════════════════════
# 六、质量基线（面试可以主动说）
# ═══════════════════════════════════════════════════════════
h1('六、全项目质量基线', red=True)
table(
    ['检查项', '结果'],
    [
        ['8 个后端 pytest', '全部通过（③28个用例最多）'],
        ['8 个前端 vue-tsc', '全部 0 错误'],
        ['16 个服务本地启动', '全部 healthy'],
        ['数据飞轮闭环', '①⑤④→⑥→⑧ 全链路实测通过'],
        ['真实 AI 问答', 'DeepSeek 流式回复实测正常'],
        ['外部图库实调', 'Unsplash/Pexels 真实图片视频返回'],
    ],
    red_cols=(0, 1),
)
body('', red=True)

doc.save('面试备战手册.docx')
print('面试备战手册.docx 已生成')
