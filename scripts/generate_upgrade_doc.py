# -*- coding: utf-8 -*-
"""基于《面试手册升级版.md》生成 Word 版。
修改处用红色标出（用户已打印纸质版，对照红笔标注）。
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = "面试手册升级版.md"
OUT = "面试手册升级版.docx"

RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x1A, 0x1A, 0x1A)

# ── 红色修改清单：old片段 → (新片段) ─────────────────────────
# 段落里出现 old 片段时，旧文字保留黑色删除线下标记，新文字红色
RED_EDITS = [
    # Q1 自我介绍：Docker 已弃用
    (
        "Docker Compose编排17个容器统一部署",
        "8个独立端口本地运行，一键启动脚本统一管理",
    ),
    # Q1 其他 Docker 痕迹
    ("Docker部署", "本地一键启动"),
    # Q5 跨项目联动末尾：Docker DNS → 内部认证
    (
        "这些联动是通过Docker网络DNS实现的——容器之间用服务名就能通信，不写死IP。共享代码放在一个shared目录里，包括数据推送客户端和跨项目动作注册表。",
        "这些联动通过 X-Internal-Call 共享密钥内部认证实现——⑥数据中心验证密钥头，①⑤④自动推送数据无需登录。跨项目调用用 Docker 服务名加 localhost 双 URL 回退。共享代码放在 shared 目录（数据推送客户端 + 动作注册表）。",
    ),
    # 附录 ②③④ 描述更新
    ("| ② | 自媒体内容助手 | 5平台文案/标题/脚本生成 | Prompt工程、多平台适配 |",
     "| ② | 自媒体内容助手 | 5平台文案/标题/脚本生成 + 多AI供应商回退 | Prompt工程、多平台适配 |"),
    ("| ③ | 短视频脚本工坊 | 分镜表+TTS语音+字幕导出 | 结构化JSON输出 |",
     "| ③ | 短视频脚本工坊 | 分镜表+TTS语音+字幕导出 + 多AI供应商回退 | 结构化JSON输出 |"),
    ("| ④ | DAM素材管理 | AI自动标签+文搜图+毛玻璃UI | 竞态修复、上传队列、安全 |",
     "| ④ | DAM素材管理 | AI自动标签+文搜图+Unsplash/Pexels外部图库 | 竞态修复、上传队列、安全 |"),
    # 关键数字：新增数据飞轮
    (
        "3层降级         # ChromaDB→内存→降级提示",
        "3层降级         # ChromaDB→内存→降级提示\n4数据源外部图库  # Picsum/Unsplash/Pexels照片/Pexels视频\n3供应商回退      # DeepSeek→智谱GLM→通义千问\n8/8测试全绿     # pytest全过 + vue-tsc零错误",
    ),
]

# ── 整章红色标注（这些章是新增内容，整体红） ─────────────────
RED_HEADINGS = {
    "第十六章：最新升级速览（2026-08-13）",
}

# ── 新增章节内容（追加到文档末尾，整体红色） ─────────────────
NEW_CHAPTER = """第十六章：最新升级速览（2026-08-13）

> 这一章是面试前新增的——把最近两周的升级浓缩成10条，每条都能展开讲。纸质版没有这章，红笔补在第15章之后。

### 升级1：多AI供应商自动回退（②③）

DeepSeek 失败自动切换智谱 GLM-4-Flash，再失败切换通义 Qwen-Turbo。三个供应商都是 OpenAI 兼容接口，统一封装在 multi_provider.py。三级优先：微调代理 → DeepSeek → 多供应商。占位符 Key 自动跳过。

面试说法："我实现了多供应商容灾回退——单一 API 故障不影响服务可用性。"

### 升级2：外部图库 API 接入（④）

Unsplash（图片）+ Pexels（照片+视频）+ Lorem Picsum（免Key）三图库。前端"🌍外部图库"弹窗4数据源切换，关键词搜索+一键导入素材库（自动AI打标）。未配Key时优雅降级到免Key图库并友好提示。

### 升级3：数据飞轮真正闭环（全项目）

X-Internal-Call 共享密钥内部认证：①客服对话、⑤培训记录、④素材标签自动推送⑥数据中心，⑥清洗标注后导出给⑧微调，微调模型部署为推理API回馈其他项目。实测闭环：⑤结束训练会话后⑥自动收到数据；⑧从⑥导入数据集创建微调任务成功。

### 升级4：⑤⑥⑦⑧功能补齐

⑥版本对比/质量报告页/数据集页/标注人工验证；⑦任务自动拆解（每Agent独立子任务）+Agent心跳；⑧RLHF偏好数据导出/任务停止/学习率曲线/模型版本对比；⑤进步曲线图表/会话只读回放。

### 升级5：安全审计修复

8个.env硬编码真实API Key→占位符+.env.example模板；管理员弱密码123456→ChangeMe!2024；④文件端点补认证；forgot-password用户枚举泄漏修复（③④）；上传扩展名白名单。

### 升级6：难点故事新增6个

跨项目推送401（X-Internal-Call）、SSE执行ID在flush前取用（4种编排模式全崩）、Windows GBK控制台emoji崩溃、密钥硬编码"true"静默401、上传标签Form注解丢失、seed_admin_user密码不同步。

### 升级7：质量基线

8个后端pytest全部通过（③28个用例最多）；8个前端vue-tsc零错误；16个服务本地启动全部healthy；真实DeepSeek流式问答实测正常。

### 升级8：Docker决策调整

本地开发不再依赖Docker（面试演示更快更稳）。若被问部署：诚实说"开发用本地uvicorn+vite，生产环境Docker Compose方案已设计并验证过，代码仓库保留部署方案"。

### 升级9：面试演示五步流

①客服真实AI提问看流式+引用 → ⑦多智能体看SSE任务拆解 → ⑥数据中心展示数据飞轮+质量报告 → ⑧微调展示从⑥导入数据训练 → ④素材展示外部图库搜图导入。

### 升级10：文档体系

CHANGELOG.md（19个难点故事）+ 面试手册升级版（本文档）+ 面试导航.html（8项目一键打开）。
"""


def add_runs(p, text, color=BLACK):
    """添加一个 run，自动处理粗体标记 **text**。"""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.color.rgb = color
    return p


def add_paragraph_with_edits(doc, text, style_normal=True):
    """段落级：应用 RED_EDITS 替换，旧片段红色删除线，新片段红色。"""
    new_text = text
    for old, new in RED_EDITS:
        if old in new_text:
            p = doc.add_paragraph()
            # 分段处理
            idx = new_text.find(old)
            add_runs(p, new_text[:idx])
            r = p.add_run(old)
            r.font.color.rgb = RED
            r.font.strike = True  # 旧文字红删除线
            add_runs(p, new, RED)
            add_runs(p, new_text[idx + len(old):])
            return
    p = doc.add_paragraph()
    add_runs(p, text)


def convert_md(src, out):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)

    lines = open(src, encoding='utf-8').read().split('\n')
    i = 0
    in_code = False
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if code_buf:
            for cl in code_buf:
                p = doc.add_paragraph()
                r = p.add_run(cl)
                r.font.name = 'Consolas'
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_buf = []

    while i < len(lines):
        line = lines[i]

        # 代码块
        if line.strip().startswith('```'):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 表格：收集连续 | 行
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                t = lines[i].strip()
                if not re.match(r'^\|[\s:\-|]+\|$', t):  # 跳过 |---| 分隔行
                    cells = [c.strip() for c in t.strip('|').split('|')]
                    table_lines.append(cells)
                i += 1
            if table_lines:
                t = doc.add_table(rows=len(table_lines), cols=len(table_lines[0]))
                t.style = 'Table Grid'
                for ri, row in enumerate(table_lines):
                    for ci in range(len(table_lines[0])):
                        val = row[ci] if ci < len(row) else ''
                        cell = t.rows[ri].cells[ci]
                        cell.text = val
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(9)
                                if ri == 0:
                                    r.bold = True
            continue

        # 标题
        if line.startswith('#### '):
            p = doc.add_paragraph()
            r = p.add_run(line[5:])
            r.bold = True
            r.font.size = Pt(11)
            i += 1
            continue
        if line.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(line[4:])
            r.bold = True
            r.font.size = Pt(11.5)
            i += 1
            continue
        if line.startswith('## '):
            text = line[3:]
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = RED if text in RED_HEADINGS else BLACK
            i += 1
            continue
        if line.startswith('# '):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(20)
            i += 1
            continue

        # 引用
        if line.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            content = line.lstrip('> ').strip()
            # 引用段落同样应用红色修改
            applied = False
            for old, new in RED_EDITS:
                if old in content:
                    idx = content.find(old)
                    add_runs(p, content[:idx])
                    r = p.add_run(old)
                    r.font.color.rgb = RED
                    r.font.strike = True
                    add_runs(p, new, RED)
                    add_runs(p, content[idx + len(old):])
                    applied = True
                    break
            if not applied:
                add_runs(p, content)
            for r in p.runs:
                r.font.size = Pt(10)
                r.italic = True
            i += 1
            continue

        # 分隔线
        if line.strip() == '---':
            i += 1
            continue

        # 普通段落
        if line.strip():
            add_paragraph_with_edits(doc, line.strip())
        i += 1

    flush_code()

    # ── 追加新章节（整体红色） ──
    doc.add_paragraph()
    nc_lines = NEW_CHAPTER.strip().split('\n')
    for l in nc_lines:
        if l.startswith('### '):
            p = doc.add_paragraph()
            r = p.add_run(l[4:])
            r.bold = True
            r.font.size = Pt(11.5)
            r.font.color.rgb = RED
        elif l.startswith('## '):
            p = doc.add_paragraph()
            r = p.add_run(l[3:])
            r.bold = True
            r.font.size = Pt(13)
            r.font.color.rgb = RED
        elif l.startswith('>'):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            add_runs(p, l.lstrip('> ').strip(), RED)
            for r in p.runs:
                r.font.size = Pt(10)
                r.italic = True
        elif l.strip():
            p = doc.add_paragraph()
            add_runs(p, l, RED)

    doc.save(out)
    print(f"{out} 已生成（基于 {src}）")


if __name__ == '__main__':
    convert_md(SRC, OUT)
